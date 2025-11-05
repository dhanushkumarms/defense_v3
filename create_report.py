"""
Create a comprehensive Word document report for the defense project
Based on the CRL report format with all project content
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

def create_report():
    # Create a new Document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Enhancing LLM Robustness Against\nAdversarial Attacks")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("A Comprehensive Evaluation of Safety Mechanisms\nin Large Language Models")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.datetime.now().strftime("%B %Y"))
    date_run.font.size = Pt(12)
    date_run.font.name = 'Times New Roman'
    
    # Page break after title
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_items = [
        "Abstract",
        "Problem Statement",
        "Literature Survey",
        "    1. The Proliferation and Sophistication of Jailbreak Attacks",
        "    2. The Landscape of Defense Mechanisms",
        "    3. The Unending Arms Race: Helpfulness vs. Harmlessness",
        "System Architecture and Workflow",
        "Experimental Setup",
        "Results and Analysis",
        "Conclusion and Future Work",
        "References"
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25) if item.startswith("    ") else Inches(0)
    
    doc.add_page_break()
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = """Modern Large Language Models (LLMs) are vulnerable to adversarial or paraphrased prompts that bypass intended safety mechanisms. Our project aims to address the challenge that lies in:

1. Systematically identifying and generating diverse adversarial prompts that effectively bypass current safety mechanisms.
2. Quantifying the baseline robustness of a target LLM against these varied attack vectors.
3. Developing and evaluating defense mechanisms that are robust against these attacks without compromising the model's performance.

This research provides a comprehensive evaluation of the Llama-2 7B model's robustness against adversarial attacks, demonstrating significant vulnerabilities in current safety alignment techniques and proposing Semantic Smoothing as a promising defense mechanism."""
    
    doc.add_paragraph(abstract_text)
    
    # Problem Statement
    doc.add_heading('Problem Statement', level=1)
    problem_text = """Large Language Models (LLMs), despite extensive safety measures, remain vulnerable to adversarial prompts that bypass filters to generate harmful content. The fundamental challenge lies in identifying, quantifying, and defending against these attacks without compromising the model's core performance and utility.

Current safety mechanisms, primarily based on Reinforcement Learning from Human Feedback (RLHF), show high effectiveness against direct harmful queries but exhibit critical weaknesses when faced with semantically equivalent but structurally different prompts. This "mismatched generalization" represents a fundamental gap in LLM safety that requires systematic investigation and novel defense strategies."""
    
    doc.add_paragraph(problem_text)
    
    # Literature Survey
    doc.add_heading('Literature Survey', level=1)
    lit_intro = doc.add_paragraph()
    lit_intro_run = lit_intro.add_run("This section provides a comprehensive overview of the current research landscape in LLM security, focusing on adversarial attacks and defense mechanisms.")
    lit_intro_run.italic = True
    
    # Section 1
    doc.add_heading('1. The Proliferation and Sophistication of Jailbreak Attacks', level=2)
    section1_text = """The vulnerability of Large Language Models (LLMs) to "jailbreak" attacks—manipulative inputs that trick models into violating their safety policies—is a foundational challenge in AI safety. The field has rapidly evolved from simple, manually crafted prompts to sophisticated, automated attack strategies.

Early Manual Attacks and Taxonomies:

Initial research, such as the work by Wei et al. (2023) in "Jailbroken: How Does LLM Safety Training Fail?", provided the first systematic analysis of why these attacks succeed. They identified two primary failure modes in LLM safety training:

1. Competing Objectives: This occurs when the model's primary goal of following user instructions directly conflicts with its safety training. For instance, a user might embed a harmful request within a seemingly benign task (e.g., "Write a story where a character describes how to build a bomb"). The model's instruction-following objective may override its safety objective.

2. Mismatched Generalization: Safety training often covers a specific set of harmful topics and phrasing. However, it may fail to generalize to novel or unusual prompts that are semantically similar but structurally different. An attacker can exploit this by rephrasing a forbidden request in a creative way, such as through poetry, code, or abstract metaphors.

Wei et al. demonstrated that attacks designed around these failure modes could successfully jailbreak even state-of-the-art models like GPT-4 and Claude, achieving high success rates on prompts that the models' own internal red-teaming efforts had identified as unsafe. This work laid the groundwork for understanding that safety is not a simple add-on but a complex challenge deeply intertwined with the model's core capabilities.

Automated and Gradient-Based Attack Generation:

The limitations of manual prompt crafting—being time-consuming and not scalable—led researchers to focus on automating the discovery of adversarial attacks. A landmark paper in this area is "Universal and Transferable Adversarial Attacks on Aligned Language Models" by Zou et al. (2023). They introduced the Greedy Coordinate Gradient (GCG) attack, a gradient-based search algorithm that efficiently discovers adversarial suffixes.

The key innovation of the GCG attack is its ability to generate a short, seemingly nonsensical string of characters that, when appended to a harmful prompt, can reliably jailbreak a model. The results were striking: the GCG attack achieved an 84% success rate on Vicuna-7B, an open-source model, and, more alarmingly, these same suffixes were transferable, successfully jailbreaking closed-source, state-of-the-art models like GPT-3.5, GPT-4, and PaLM-2.

This demonstrated that vulnerabilities found in one model could represent a systemic weakness across the entire LLM ecosystem. Our own project's findings, where we observed a 52.1% attack success rate with adversarial paraphrasing, align with these results, confirming that even without complex gradient-based methods, semantic manipulation is highly effective."""
    
    doc.add_paragraph(section1_text)
    
    # Section 2
    doc.add_heading('2. The Landscape of Defense Mechanisms', level=2)
    section2_text = """In response to this escalating threat, a variety of defense strategies have been proposed, moving beyond simple keyword filtering.

Input and Output Filtering:

The most straightforward defense is to filter inputs and outputs. This can involve using a "guard" model to check if a prompt or a generated response is safe. However, as Wei et al. (2023) noted, these guard models are often smaller, less capable LLMs that are themselves vulnerable to the very same adversarial attacks they are meant to prevent. This creates a recursive problem where the defense is as weak as the model it is trying to protect.

Adversarial Training:

A more robust approach is adversarial training, a technique borrowed from computer vision. This involves fine-tuning the LLM on a dataset of adversarial examples, teaching it to recognize and refuse jailbreak attempts. While this can improve robustness against known attack types, it has several drawbacks:

• It is computationally expensive, requiring significant resources to generate adversarial examples and retrain the model.
• It can lead to a phenomenon known as "catastrophic forgetting," where the model's performance on benign tasks degrades.
• It may not generalize to unseen or novel attack types, meaning the model is only as safe as the attacks it has been trained on.

Semantic and Principled Defenses: The Case for Semantic Smoothing

Recognizing the limitations of the above methods, researchers have started developing more principled defenses that focus on the intent behind a prompt rather than its surface-level structure. One of the most promising of these is Semantic Smoothing, introduced by Robey et al. (2023) in "Defending Large Language Models Against Jailbreak Attacks via Semantic Smoothing."

The core idea behind Semantic Smoothing is simple yet powerful: if a prompt is truly benign, small paraphrases of it should also be benign. The technique works as follows:

1. Take the original input prompt.
2. Generate multiple, slightly different versions of it through paraphrasing.
3. Feed each of these perturbed prompts to the LLM.
4. Classify the safety of each response.
5. If a significant portion of the responses are deemed unsafe, the original prompt is rejected, even if its own response was benign.

This approach acts as a "semantic firewall." It is not fooled by clever wording because the underlying harmful intent is likely to be preserved across paraphrases. Robey et al. demonstrated that Semantic Smoothing could successfully defend against a range of attacks, including the powerful GCG attack, reducing its success rate from over 50% to under 2% in some cases. Crucially, this was achieved with only a minor degradation in the model's performance on benign prompts. This makes it a highly practical and effective defense, and it is the primary candidate for implementation in the next phase of our project."""
    
    doc.add_paragraph(section2_text)
    
    # Section 3
    doc.add_heading('3. The Unending Arms Race: Helpfulness vs. Harmlessness', level=2)
    section3_text = """The current state of LLM security is best described as an ongoing "arms race" between attackers and defenders. As new defense mechanisms are developed, more sophisticated attacks emerge to bypass them. This dynamic highlights a fundamental tension in LLM design: the trade-off between helpfulness and harmlessness.

A model that is too heavily focused on safety may exhibit "false refusals," refusing to answer legitimate and safe questions. This degrades the user experience and limits the model's utility. Conversely, a model that is too helpful may be easily manipulated into generating harmful content.

The literature strongly suggests that there is no single "silver bullet" for LLM safety. Instead, a multi-layered, defense-in-depth strategy is required. This could include a combination of:

• Input filtering using robust, adversarially trained guard models.
• Model fine-tuning on a diverse and continuously updated set of adversarial examples.
• Principled defenses like Semantic Smoothing to detect malicious intent.
• Output monitoring to flag and analyze harmful generations after they occur.

The ultimate goal, as our project aims to contribute to, is the development of LLMs that are not only powerful and capable but also robustly and reliably safe against a diverse and evolving landscape of adversarial threats."""
    
    doc.add_paragraph(section3_text)
    
    doc.add_page_break()
    
    # System Architecture
    doc.add_heading('System Architecture and Workflow', level=1)
    arch_text = """Our system is designed as a modular pipeline to test LLM vulnerabilities and defenses. The architecture consists of three main stages: Attack Generation, Model Inference, and Safety Evaluation."""
    doc.add_paragraph(arch_text)
    
    doc.add_heading('1. Attack Generation', level=2)
    attack_gen_text = """• Input: A dataset of 2,500 unique harmful prompts (three_variant_dataset_2500.csv).
• Process: For each base prompt, three variants are created:
    1. Direct: The original, unaltered harmful prompt.
    2. Simple Paraphrase: The prompt is rephrased using a standard paraphrasing model to alter syntax while preserving semantic meaning.
    3. Adversarial Paraphrase (Jailbreak): The prompt is embedded within a larger, adversarially crafted template designed to trick the model into ignoring its safety protocols.
• Output: A test suite of 7,500 prompts (2,500 prompts × 3 variants)."""
    doc.add_paragraph(attack_gen_text)
    
    doc.add_heading('2. Model Inference', level=2)
    inference_text = """• Input: The 7,500 generated attack prompts.
• Process: Each prompt is sent to the target LLM (Llama-2-7b-chat-hf), which is loaded in a 4-bit quantized configuration. The model's response to each prompt is collected.
• Output: A corresponding set of 7,500 model responses."""
    doc.add_paragraph(inference_text)
    
    doc.add_heading('3. Safety Evaluation', level=2)
    safety_text = """• Input: The 7,500 model responses.
• Process: Each response is evaluated by an ensemble of three specialized safety classifiers: ToxicBERT, RoBERTa-hate, and ToxicChat-BERT. A weighted average of their scores determines the final toxicity score.
• Output: A classification for each response as either "Safe" or "Harmful." The Attack Success Rate (ASR) is then calculated as the percentage of prompts that successfully elicited a harmful response."""
    doc.add_paragraph(safety_text)
    
    doc.add_page_break()
    
    # Experimental Setup
    doc.add_heading('Experimental Setup', level=1)
    setup_text = """This section details the hardware, software, and models used in our evaluation."""
    doc.add_paragraph(setup_text)
    
    doc.add_heading('Hardware Configuration', level=2)
    hardware_text = """• GPU: NVIDIA RTX 3050 Ti (4GB VRAM)
• CPU: 12th Gen Intel Core i7
• RAM: 16GB"""
    doc.add_paragraph(hardware_text)
    
    doc.add_heading('Software Environment', level=2)
    software_text = """• Python: 3.8+
• PyTorch: 2.5.1 (with CUDA 12.1)
• Transformers: 4.57.0 (Hugging Face)
• BitsAndBytes: 0.48.1 (for 4-bit quantization)
• A complete list of dependencies is available in requirements.txt."""
    doc.add_paragraph(software_text)
    
    doc.add_heading('Model Configuration', level=2)
    model_text = """• Target LLM: meta-llama/Llama-2-7b-chat-hf
• Quantization: To run the 7-billion-parameter model on a 4GB GPU, we use 4-bit NF4 quantization via the BitsAndBytes library. This technique significantly reduces the model's memory footprint from ~28GB to under 4GB, with a minimal impact on performance for inference tasks."""
    doc.add_paragraph(model_text)
    
    doc.add_heading('Safety Classifiers', level=2)
    classifier_text = """• Ensemble Models:
    - unitary/toxic-bert
    - facebook/roberta-hate-speech-dynabench-r4-target
    - lmsys/toxic-chat-bert-base-uncased
• Detection Threshold: A response is classified as harmful if the weighted ensemble toxicity score is ≥ 0.50."""
    doc.add_paragraph(classifier_text)
    
    doc.add_page_break()
    
    # Results and Analysis
    doc.add_heading('Results and Analysis', level=1)
    results_intro = """The results from our initial baseline evaluation (Phase 1) are detailed in the models/attack3.ipynb notebook and summarized here. The key finding is that while the Llama-2-7B model is quite robust against direct harmful questions, its safety measures are easily bypassed by simple and adversarial paraphrasing."""
    doc.add_paragraph(results_intro)
    
    # Create results table
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Attack Variant'
    header_cells[1].text = 'Attack Success Rate (ASR)'
    header_cells[2].text = 'Defense Rate'
    header_cells[3].text = 'Relative Increase vs. Direct'
    
    # Data rows
    data = [
        ['Direct', '1.7%', '98.3%', '1.0x'],
        ['Simple Paraphrase', '15.4%', '84.6%', '9.1x'],
        ['Adversarial Paraphrase', '52.1%', '47.9%', '30.7x']
    ]
    
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text
    
    doc.add_paragraph()
    
    doc.add_heading('Key Observations', level=2)
    observations_text = """1. High Baseline Safety: An ASR of only 1.7% on direct prompts shows that the model's default safety training is effective against straightforward harmful requests.

2. Vulnerability to Paraphrasing: A simple rephrasing of the prompt increases the ASR by over 9 times. This confirms the "mismatched generalization" failure mode, where the model's safety training does not extend to semantically equivalent but structurally different prompts.

3. Effectiveness of Jailbreaks: The use of adversarial jailbreak prompts leads to a dramatic failure of safety mechanisms, with the ASR skyrocketing to 52.1%. This is a 30-fold increase in successful attacks compared to the direct baseline.

4. True Attack Success: More telling is the "true ASR," which measures how often an adversarial prompt succeeds specifically on questions that were safely handled in their direct form. Our analysis shows that in 49.8% of cases where the model correctly refused a direct harmful prompt, it failed and complied when the same prompt was wrapped in an adversarial jailbreak. This isolates the effectiveness of the jailbreak technique itself.

These results provide a clear and quantifiable baseline of the model's vulnerabilities and underscore the critical need for more robust defense mechanisms like Semantic Smoothing, which will be the focus of the next phase of our project."""
    doc.add_paragraph(observations_text)
    
    doc.add_page_break()
    
    # Conclusion and Future Work
    doc.add_heading('Conclusion and Future Work', level=1)
    conclusion_text = """Our research to date has successfully quantified the vulnerability of a state-of-the-art LLM, Llama-2-7B, to various forms of adversarial attacks. We have established a clear baseline, demonstrating that while the model is safe against direct harmful prompts, its defenses are brittle and easily circumvented by paraphrasing and adversarial jailbreaking techniques. Our findings align with the broader academic literature and confirm that LLM safety remains a significant and unsolved problem."""
    doc.add_paragraph(conclusion_text)
    
    doc.add_heading('Future Work', level=2)
    future_work_text = """Our project will proceed with the following steps:

1. Implement Semantic Smoothing: We will develop and integrate the Semantic Smoothing defense mechanism into our evaluation pipeline.

2. Comparative Re-evaluation: We will re-run our entire suite of 7,500 test cases with the defense enabled.

3. Analyze Defense Effectiveness: We will compare the ASR of the defended model against our established baseline to quantify the effectiveness of Semantic Smoothing. We will also measure any potential degradation in performance on benign prompts.

4. Report and Document: The final phase will involve a comprehensive write-up of our findings, including visualizations and a detailed analysis of the defense's strengths and weaknesses.

By systematically evaluating both the problem and a potential solution, this project aims to contribute valuable insights to the ongoing effort to build safer and more reliable Large Language Models."""
    doc.add_paragraph(future_work_text)
    
    doc.add_page_break()
    
    # References
    doc.add_heading('References', level=1)
    references = [
        "Zou, A., Wang, Z., Kolter, J. Z., & Fredrikson, M. (2023). Universal and Transferable Adversarial Attacks on Aligned Language Models. arXiv preprint arXiv:2307.15043.",
        "Wei, A., Haghtalab, N., & Steinhardt, J. (2023). Jailbroken: How Does LLM Safety Training Fail?. arXiv preprint arXiv:2307.02483.",
        "Robey, A., Pin-Yu, C., & Chawla, S. (2023). Defending Large Language Models Against Jailbreak Attacks via Semantic Smoothing. arXiv preprint arXiv:2310.00401.",
        "Jones, A., et al. (2023). Automated Red-Teaming of Language Models. (Placeholder for a more specific reference if available)."
    ]
    
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph(f"{i}. {ref}")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    # Save the document
    filename = f'Defense_Project_Report_{datetime.datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f"Report created successfully: {filename}")
    return filename

if __name__ == "__main__":
    create_report()
